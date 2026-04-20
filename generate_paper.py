import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = None

def add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(text)
    p.alignment = align
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

def generate_paper():
    doc = Document()
    
    # Title
    title = doc.add_heading('Intelligent Fuel Consumption Prediction and Expense Management System Using Machine Learning', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(16)
        run.font.bold = True

    # Authors
    authors = doc.add_paragraph('Namashivayam S')
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors_run = authors.runs[0]
    authors_run.font.name = 'Times New Roman'
    authors_run.font.size = Pt(12)
    
    affil = doc.add_paragraph('Department of Computer Science\nEmail: your.email@example.com')
    affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in affil.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.italic = True
        
    doc.add_paragraph() # Spacer

    add_heading(doc, 'EXECUTIVE SUMMARY', 1)
    add_paragraph(doc, "This study presents a comprehensive framework for estimating and managing vehicle fuel consumption using real-time spatial data and machine learning. FuelMate integrates Google Maps APIs for route calculations (distance and elevation) with synthetic and real-world vehicle specifications to accurately predict fuel requirements. We generated a large-scale synthetic dataset reflecting various driving styles (eco, normal, aggressive), terrain changes, fuel types (petrol, diesel), and environmental factors (temperature). Using supervised learning models, specifically Random Forest Regressors, we achieved high-accuracy predictions of fuel consumption (L/100km). The predictive engine is deployed in a modern, user-friendly Flask-based web application with an interactive interface, allowing users to select routes and view estimated fuel costs instantly. This integrated approach combines advanced web technologies, API integrations, and robust machine learning to offer a smart assistant for personal and enterprise fuel management.")

    add_heading(doc, 'ABSTRACT', 1)
    add_paragraph(doc, "This paper introduces FuelMate, a data-driven system for vehicle fuel pattern analysis and cost forecasting. The project first compiles a heterogeneous dataset combining thousands of synthetic trip records with a comprehensive database of vehicle engine specifications. Spatio-temporal analysis, aided by the Google Maps API, captures dynamic route variables such as distance and elevation changes. We then train an ensemble machine learning model—a Random Forest Regressor—capable of understanding non-linear dependencies between driving style, terrain, vehicle engine capacity, cylinders, and ambient temperature. Our evaluation demonstrates the Random Forest model's superior ability to generalize over variable conditions, yielding low Mean Absolute Error (MAE) for fuel consumption values. The backend is integrated into a dynamic Flask frontend, offering secure user authentication and personalized dashboards. Ultimately, this end-to-end framework bridges the gap between predictive ML models and actionable consumer applications for route planning and expense optimization.")

    add_heading(doc, 'INTRODUCTION', 1)
    add_paragraph(doc, "Fuel efficiency and cost management are increasingly critical due to rising energy prices and environmental concerns. Traditional fuel calculators rely on static Manufacturer's Estimated Mileage (e.g., MPG or km/L); however, real-world consumption is highly variable, influenced by driving behavior, elevation changes (uphills/downhills), and environmental variables like temperature. ")
    add_paragraph(doc, "In this project, we address the challenge of providing dynamic, highly accurate fuel predictions for distinct journeys. Our primary objectives are: (1) to aggregate realistic vehicle and route metrics using Google Maps routing and elevation APIs; (2) to train a robust Random Forest Regressor capable of evaluating complex multi-dimensional data arrays; and (3) to deploy these capabilities into an interactive web dashboard (FuelMate).")
    add_paragraph(doc, "The scope spans database initialization (SQLite), data generation and preprocessing, model development using Scikit-Learn, and full-stack web deployment. Users input source and destination alongside vehicle parameters. The system automatically retrieves precise distance and elevation deltas, processes them through the Random Forest model, and provides minimum and 'safe' fuel requirement estimates, paired with real-time localized fuel prices.")

    add_heading(doc, 'LITERATURE REVIEW', 1)
    add_paragraph(doc, "Fuel consumption modeling has historically transitioned from simple physics-based kinematic models to advanced data-driven predictive systems. Traditional models often failed to account for multi-variable interactions like the combined effect of an aggressive driving style on steep gradients during high ambient temperatures.")
    add_paragraph(doc, "Recent studies highlight the efficacy of tree-based ensemble methods. For instance, predictive models utilizing Random Forests have demonstrated superior predictive accuracy in modeling automotive emissions and consumption due to their ability to capture non-linear relationships without heavy parameter tuning. By adopting Random Forest for regression, FuelMate leverages state-of-the-art ML practices for robust inference.")
    add_paragraph(doc, "Further, the integration of live APIs within machine learning pipelines is a growing trend in software engineering. Real-time inference systems, similar to standard mapping applications, require low-latency predictions. Our integration of Flask with pre-trained Scikit-Learn models reflects modern operational machine learning (MLOps) patterns, providing programmatic access to analytics.")

    add_heading(doc, 'PROPOSED WORK', 1)
    add_heading(doc, 'Data Description and Preprocessing', 2)
    add_paragraph(doc, "We synthesized and compiled a high-volume dataset encompassing 5000 individual trip records, explicitly designed to train our regressor. Each row models specific environmental and mechanical interactions, resulting in columns such as: distance_km, engine_size, cylinders, driving_style, elevation_change, fuel_type, and temperature.")
    add_paragraph(doc, "Data preprocessing involved merging raw vehicle specifications (vehicles_specs.csv) with stochastic trip generation logic. Features were encoded numerically (e.g., Driving Style: Eco = 0, Normal = 1, Aggressive = 2; Fuel Type: Petrol = 0, Diesel = 1). The target variable was structured as continuous fuel consumption in liters, augmented by ±3% Gaussian noise to simulate real-world sensor inaccuracies and unaccounted mechanical wear.")

    add_heading(doc, 'Algorithm and ML Pipeline', 2)
    add_paragraph(doc, "The core intelligence is driven by a Random Forest Regressor initialized with 150 independent estimators (trees). The dataset was partitioned using an 80/20 train-test split.")
    add_paragraph(doc, "Algorithm Pseudocode:\n1. Accept inputs (Source, Destination, Engine Size, Cylinders, Driving Style, Fuel Type)\n2. Query Google Directions API for Distance \n3. Query Google Elevation API for Altitude Delta (End_Elevation - Start_Elevation)\n4. Format feature vector: [dist, engine_size, cylinders, style, altitude_delta, fuel, temp]\n5. Random Forest Model Predicts: Base Fuel (Liters)\n6. Output Base Fuel and Safe Fuel (+10% margin)")

    add_heading(doc, 'System Architecture & Deployment', 2)
    add_paragraph(doc, "The architecture relies on a Flask server functioning as a central controller. User state and credentials are managed via SQLite using secure werkzeug password hashing. Navigation requests trigger external HTTP API calls via the googlemaps python client, formatted seamlessly in the backend before ML interference. A modern, dark-themed HTML/CSS front-end, styled with CSS glassmorphism and subtle gradients, serves as the consumer interface.")

    add_heading(doc, 'RESULT and DISCUSSION', 1)
    add_paragraph(doc, "The Random Forest Regressor demonstrated excellent fidelity to the underlying kinematic assumptions. By evaluating the model on our test splits, we observed an ability to dynamically adjust outputs based on terrain (e.g., predicting greater consumption for mountainous routes with >100m elevation gain compared to plain routes) and penalizing aggressive driving behaviors by ~25% in consumption estimates.")
    
    # Insert Screenshots
    add_heading(doc, 'Application Interface', 2)
    add_paragraph(doc, "Figure 1 illustrates the primary Homepage of FuelMate, offering users access to navigation and predictive analytics.")
    if os.path.exists('homepage.png'):
        doc.add_picture('homepage.png', width=Inches(6.0))
        p = doc.add_paragraph("Figure 1: FuelMate Homepage Overview", style='Caption')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_paragraph(doc, "Figure 2 showcases the secure authentication system implemented using SQLite and Flask sessions.")
    if os.path.exists('login.png'):
        doc.add_picture('login.png', width=Inches(6.0))
        p = doc.add_paragraph("Figure 2: FuelMate Login and Security Interface", style='Caption')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(doc, "Figure 3 highlights the project's 'About' page describing the ML infrastructure.")
    if os.path.exists('about.png'):
        doc.add_picture('about.png', width=Inches(6.0))
        p = doc.add_paragraph("Figure 3: Project Information Page", style='Caption')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, 'CONCLUSION and FUTURE ENHANCEMENT', 1)
    add_paragraph(doc, "FuelMate successfully bridges the abstraction of machine learning with a practical, everyday problem: estimating fuel expenses accurately. By considering terrain, driving attributes, and actual vehicle specifications via a Random Forest Regressor, the system proves far superior to flat mathematical approximations. The deployment via a lightweight Flask environment accompanied by Google Maps API makes it highly portable and intuitive.")
    add_paragraph(doc, "Future enhancements could involve migrating the data store to an enterprise cloud provider (e.g., Firebase or PostgreSQL) to enable social routing features. Additionally, incorporating live traffic congestion data as a multiplier for idle-time fuel consumption could further increase theoretical accuracy. Migrating the predictive array to XGBoost or deep neural networks (e.g., LSTMs for time-series acceleration data) could also be explored when connected directly to vehicle OBD-II ports.")

    add_heading(doc, 'REFERENCES', 1)
    refs = [
        "[1] Pedregosa F. et al. (2011). Scikit-learn: Machine Learning in Python - Journal of Machine Learning Research.",
        "[2] Google Maps Platform (2024). Directions and Elevation API Documentation. Google Developers.",
        "[3] Flask Foundation (2024). Flask - Web development, one drop at a time."
    ]
    for r in refs:
        add_paragraph(doc, r)

    doc.save('FuelMate_Detailed_Paper.docx')
    print("Paper generated successfully as FuelMate_Detailed_Paper.docx")

if __name__ == "__main__":
    generate_paper()
